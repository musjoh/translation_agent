from __future__ import annotations

import hmac
import os
from pathlib import Path
import re
import tempfile
from typing import Any

import streamlit as st
from docx import Document

from src.config import AppConfig
from src.pipeline import get_extraction_preview, run_translation_pipeline


def _is_streamlit_cloud() -> bool:
    return bool(os.getenv("STREAMLIT_SHARING_MODE") or os.getenv("IS_STREAMLIT_CLOUD"))


def _privacy_lock_enabled(is_cloud: bool) -> bool:
    # In cloud, privacy lock is ON by default. Set PRIVACY_LOCK_MODE=0 to disable.
    if not is_cloud:
        return False
    raw = os.getenv("PRIVACY_LOCK_MODE", "1").strip().lower()
    return raw not in {"0", "false", "off", "no"}


def _redact_value(value: str) -> str:
    value = re.sub(r"sk-[A-Za-z0-9_\-]{12,}", "[REDACTED_API_KEY]", value)
    value = re.sub(r"Bearer\s+[A-Za-z0-9_\-\.]+", "Bearer [REDACTED_TOKEN]", value, flags=re.IGNORECASE)
    return value


def _sanitize_for_display(data: Any) -> Any:
    if isinstance(data, dict):
        return {k: _sanitize_for_display(v) for k, v in data.items()}
    if isinstance(data, list):
        return [_sanitize_for_display(v) for v in data]
    if isinstance(data, str):
        return _redact_value(data)
    return data


def _require_access_password() -> bool:
    """
    Optional app-level password gate for shared access (e.g. Cloudflare Tunnel).
    Set APP_PASSWORD in environment to enable.
    """
    required = os.getenv("APP_PASSWORD", "").strip()
    if not required:
        return True

    if st.session_state.get("auth_ok") is True:
        return True

    st.subheader("Access Protected")
    st.caption("This app requires a password.")
    with st.form("access_password_form"):
        entered = st.text_input("Access password", type="password")
        submit = st.form_submit_button("Enter")

    if submit and hmac.compare_digest(entered.strip(), required):
        st.session_state["auth_ok"] = True
        st.rerun()
    elif submit:
        st.error("Invalid password.")
    return False


def _read_docx_preview(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables_preview: list[list[list[str]]] = []
    for table in doc.tables[:3]:
        rows: list[list[str]] = []
        for row in table.rows[:6]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_preview.append(rows)
    return {
        "paragraphs": paragraphs[:20],
        "tables_preview": tables_preview,
    }


def main() -> None:
    st.set_page_config(page_title="Local Document Translator", layout="wide")
    st.title("Local Document Translator")
    st.caption("MVP scaffold: PDF, DOCX, TXT -> translated DOCX")
    if not _require_access_password():
        return

    is_cloud = _is_streamlit_cloud()
    privacy_lock = _privacy_lock_enabled(is_cloud)
    if is_cloud:
        st.info("Streamlit Cloud mode detected. BYOK recommended.")
    if privacy_lock:
        st.success(
            "Privacy lock is ON: server disk persistence and checkpoint resume are disabled for this session."
        )

    with st.form("translate_form"):
        uploaded = st.file_uploader(
            "Upload document",
            type=["pdf", "docx", "txt"],
            help="Machine-readable PDF, DOCX, or TXT.",
        )
        api_key = st.text_input(
            "API Key (BYOK)",
            type="password",
            help="Used only for current requests and not stored by the app.",
        )
        base_url = st.text_input(
            "Base URL",
            value="https://api.openai.com/v1",
            help="Any OpenAI-compatible endpoint.",
        )
        model_name = st.text_input("Model", value="gpt-4o-mini")
        source_lang = st.text_input("Source Language", value="auto")
        target_lang = st.text_input("Target Language", value="zh")
        bilingual = st.checkbox("Bilingual output", value=False)
        use_mock = st.checkbox("Use mock translator (local test)", value=not is_cloud)
        if privacy_lock:
            persist_outputs = False
            resume_from_checkpoint = False
        else:
            persist_outputs = st.checkbox(
                "Save output/checkpoints on server disk",
                value=not is_cloud,
                help="Disable this on public deployments for stronger privacy.",
            )
            resume_from_checkpoint = st.checkbox(
                "Resume from last saved progress",
                value=not is_cloud,
                help="If a long run fails, next run continues from completed chunks.",
            )
        submitted = st.form_submit_button("Translate")

    if use_mock:
        st.warning(
            "Mock translator is ON. This mode is for pipeline testing and does not call the real API."
        )

    if uploaded is None:
        if submitted:
            st.error("Please upload a document.")
        return

    suffix = Path(uploaded.name).suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt"}:
        st.error("Only .pdf, .docx, and .txt are supported.")
        return

    with tempfile.TemporaryDirectory(prefix="preview_job_") as preview_tmpdir:
        preview_path = Path(preview_tmpdir) / uploaded.name
        preview_path.write_bytes(uploaded.getbuffer())
        with st.spinner("Extracting preview..."):
            preview = get_extraction_preview(preview_path)

    st.subheader("Extraction Preview")
    if "error" in preview:
        st.error(preview["error"])
    else:
        if preview.get("used_ocr_fallback"):
            st.info("OCR fallback was used for this PDF (image/scanned content detected).")
            if preview.get("ocr_debug"):
                st.json(preview["ocr_debug"])

        c1, c2 = st.columns(2)
        c1.metric("Text blocks", preview["text_block_count"])
        c2.metric("Table blocks", preview["table_block_count"])

        with st.expander("Text preview", expanded=True):
            if preview["text_preview"]:
                for idx, item in enumerate(preview["text_preview"], start=1):
                    st.markdown(
                        f"**{idx}. Page {item['page_number']} (order {item['order']})**"
                    )
                    st.text(item["snippet"])
            else:
                st.caption("No extracted text blocks.")

        with st.expander("Table preview", expanded=True):
            if preview["table_preview"]:
                for table in preview["table_preview"]:
                    st.markdown(
                        f"**{table['table_id']} (page {table['page_number']})**"
                    )
                    if table["headers"]:
                        st.write("Headers:", table["headers"])
                    if table["rows_preview"]:
                        st.table(table["rows_preview"])
                    else:
                        st.caption("No rows extracted.")
            else:
                st.caption("No extracted tables.")

    if not submitted:
        return

    if not use_mock and not api_key.strip():
        st.error("API key is required when mock translator is disabled.")
        return

    with tempfile.TemporaryDirectory(prefix="translation_job_") as tmpdir:
        input_path = Path(tmpdir) / uploaded.name
        input_path.write_bytes(uploaded.getbuffer())

        if persist_outputs:
            output_dir = Path("outputs")
            output_dir.mkdir(parents=True, exist_ok=True)
        else:
            output_dir = Path(tmpdir) / "outputs"
            output_dir.mkdir(parents=True, exist_ok=True)
            st.caption("Disk persistence disabled: outputs/checkpoints will be deleted after this run.")

        config = AppConfig(
            api_key=api_key.strip(),
            base_url=base_url.strip(),
            model_name=model_name.strip(),
            source_lang=source_lang.strip(),
            target_lang=target_lang.strip(),
            bilingual=bilingual,
            use_mock_translator=use_mock,
            resume_from_checkpoint=resume_from_checkpoint and persist_outputs,
        )

        progress_status = st.empty()
        progress_bar = st.progress(0)
        progress_log_box = st.empty()
        progress_lines: list[str] = []

        def on_progress(event: dict) -> None:
            event_type = str(event.get("type", ""))
            total_chunks = int(event.get("total_chunks", 0)) or 1
            completed_chunks = int(event.get("completed_chunks", 0))

            if event_type == "start":
                progress_status.info(
                    f"Starting translation: {completed_chunks}/{total_chunks} chunks completed."
                )
                progress_bar.progress(min(1.0, completed_chunks / total_chunks))
                reused = completed_chunks
                if reused > 0:
                    progress_lines.append(f"Resumed from checkpoint: reused {reused} chunks.")
                progress_log_box.code("\n".join(progress_lines[-20:]) or "Waiting for first chunk...")
                return

            if event_type == "chunk_success":
                chunk_id = str(event.get("chunk_id", ""))
                page_number = int(event.get("page_number", 0))
                progress_status.info(
                    f"Translating... {completed_chunks}/{total_chunks} chunks done."
                )
                progress_bar.progress(min(1.0, completed_chunks / total_chunks))
                progress_lines.append(f"[OK] page {page_number} - {chunk_id}")
                progress_log_box.code("\n".join(progress_lines[-20:]))
                return

            if event_type == "chunk_failed":
                chunk_id = str(event.get("chunk_id", ""))
                error = _redact_value(str(event.get("error", "Unknown error")))
                progress_status.error(f"Stopped at chunk {chunk_id}.")
                progress_bar.progress(min(1.0, completed_chunks / total_chunks))
                progress_lines.append(f"[FAIL] {chunk_id} - {error}")
                progress_log_box.code("\n".join(progress_lines[-20:]))
                return

            if event_type == "preflight_failed":
                error = _redact_value(str(event.get("error", "Unknown error")))
                progress_status.error("Preflight failed before translation started.")
                progress_lines.append(f"[PREFLIGHT FAIL] {error}")
                progress_log_box.code("\n".join(progress_lines[-20:]))
                return

            if event_type == "completed":
                progress_status.success(
                    f"Completed: {completed_chunks}/{total_chunks} chunks translated."
                )
                progress_bar.progress(1.0)
                progress_lines.append("[DONE] Translation completed.")
                progress_log_box.code("\n".join(progress_lines[-20:]))

        with st.spinner("Processing and translating..."):
            result = run_translation_pipeline(
                input_path=input_path,
                output_dir=output_dir,
                config=config,
                progress_callback=on_progress,
            )

        if not result.success:
            st.error(_redact_value(result.message))
            if result.debug:
                st.json(_sanitize_for_display(result.debug))
            return

        st.success(_redact_value(result.message))
        st.write(f"Output: `{result.output_path}`")
        if result.debug:
            st.json(_sanitize_for_display(result.debug))

        output_file = Path(result.output_path)
        if output_file.exists():
            preview = _read_docx_preview(output_file)
            st.subheader("Translated Result Preview")
            with st.expander("Translated text", expanded=True):
                if preview["paragraphs"]:
                    for idx, para in enumerate(preview["paragraphs"], start=1):
                        st.markdown(f"**{idx}.** {para}")
                else:
                    st.caption("No translated paragraphs found.")

            with st.expander("Translated tables", expanded=True):
                if preview["tables_preview"]:
                    for idx, rows in enumerate(preview["tables_preview"], start=1):
                        st.markdown(f"**Table {idx}**")
                        if rows:
                            st.table(rows)
                        else:
                            st.caption("Empty table.")
                else:
                    st.caption("No translated tables found.")

            with output_file.open("rb") as f:
                st.download_button(
                    "Download translated DOCX",
                    data=f.read(),
                    file_name=output_file.name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


if __name__ == "__main__":
    main()
