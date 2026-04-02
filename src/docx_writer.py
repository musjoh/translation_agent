from __future__ import annotations

from pathlib import Path

from docx import Document

from src.models import ReconstructedTable, ReconstructedText


def write_docx(
    output_path: Path,
    text_items: list[ReconstructedText],
    table_items: list[ReconstructedTable],
    bilingual: bool = False,
) -> None:
    doc = Document()
    doc.add_heading("Translated Document", level=1)

    for item in text_items:
        if bilingual and item.source_text:
            doc.add_paragraph(item.source_text)
        if item.translated_text:
            doc.add_paragraph(item.translated_text)

    for table_block in table_items:
        if bilingual and table_block.source_rows:
            doc.add_paragraph(f"Source table: {table_block.table_id}")
            _append_table(doc, table_block.source_headers, table_block.source_rows)

        doc.add_paragraph(f"Translated table: {table_block.table_id}")
        _append_table(doc, table_block.translated_headers, table_block.translated_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _append_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    num_cols = len(headers) if headers else max((len(r) for r in rows), default=1)
    table = doc.add_table(rows=1, cols=max(1, num_cols))
    table.style = "Table Grid"

    if headers:
        hdr_cells = table.rows[0].cells
        for idx, val in enumerate(headers):
            hdr_cells[idx].text = val
    else:
        table.rows[0].cells[0].text = ""

    for row in rows:
        cells = table.add_row().cells
        for idx in range(min(len(cells), len(row))):
            cells[idx].text = row[idx]
